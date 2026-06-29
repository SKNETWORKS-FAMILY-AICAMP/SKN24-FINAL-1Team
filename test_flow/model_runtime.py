from __future__ import annotations

import base64
import gc
import io
import json
import mimetypes
import os
import re
import tempfile
import time
from dataclasses import dataclass
from typing import Any

from config import (
    EMBEDDING_BATCH_SIZE,
    EMBEDDING_DIMENSIONS_INT,
    EMBEDDING_MODEL_ID,
    MAX_NEW_TOKENS,
    OCR_MAX_PDF_PAGES,
    OCR_MIN_DIRECT_TEXT_CHARS,
    OCR_MODEL_ID,
    OCR_PDF_TEXT_FIRST,
    OCR_RENDER_SCALE,
    OPENAI_API_KEY,
    OPENAI_BASE_URL,
    OPENAI_TIMEOUT_SEC,
    STT_COMPUTE_TYPE,
    STT_DEVICE,
    STT_MODEL_ID,
    TEXT_MODEL_ID,
)


OPENAI_EMBED_MAX_RETRIES = int(os.getenv("OPENAI_EMBED_MAX_RETRIES", "5"))
OPENAI_RETRY_BASE_SEC = float(os.getenv("OPENAI_RETRY_BASE_SEC", "2"))


@dataclass
class TextBundle:
    llm: Any
    backend: str


@dataclass
class EmbeddingBundle:
    tokenizer: Any
    model: Any
    device: str
    dimensions: int


@dataclass
class SttBundle:
    model: Any
    device: str
    compute_type: str


text_bundle: TextBundle | None = None
embedding_bundle: EmbeddingBundle | None = None
stt_bundle: SttBundle | None = None


def cleanup_cuda() -> None:
    gc.collect()


def _client() -> Any:
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY is not configured.")
    from openai import OpenAI

    kwargs: dict[str, Any] = {"api_key": OPENAI_API_KEY, "timeout": OPENAI_TIMEOUT_SEC}
    if OPENAI_BASE_URL:
        kwargs["base_url"] = OPENAI_BASE_URL
    return OpenAI(**kwargs)


def load_text_model() -> TextBundle:
    global text_bundle
    if text_bundle is None:
        text_bundle = TextBundle(llm=_client(), backend="openai")
    return text_bundle


def load_embedding_model() -> EmbeddingBundle:
    global embedding_bundle
    if embedding_bundle is not None:
        return embedding_bundle
    sample = embed_texts(["dimension probe"])
    embedding_bundle = EmbeddingBundle(
        tokenizer=None,
        model=EMBEDDING_MODEL_ID,
        device="api",
        dimensions=len(sample[0]),
    )
    return embedding_bundle


def load_stt_model() -> SttBundle:
    global stt_bundle
    if stt_bundle is None:
        stt_bundle = SttBundle(model=STT_MODEL_ID, device=STT_DEVICE, compute_type=STT_COMPUTE_TYPE)
    return stt_bundle


def generate_text(messages: list[dict[str, str]], *, max_new_tokens: int = MAX_NEW_TOKENS) -> str:
    bundle = load_text_model()
    response = bundle.llm.chat.completions.create(
        model=TEXT_MODEL_ID,
        messages=messages,
        temperature=0,
        max_tokens=max_new_tokens,
    )
    return str(response.choices[0].message.content or "").strip()


def extract_json(text: str) -> Any:
    cleaned = re.sub(r"<think>.*?</think>", "", str(text or ""), flags=re.DOTALL).strip()
    cleaned = re.sub(r"```(?:json)?|```", "", cleaned, flags=re.IGNORECASE).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        if not match:
            raise ValueError(f"Model did not return JSON: {cleaned[:500]}")
        return json.loads(match.group())


def generate_json(messages: list[dict[str, str]], *, max_new_tokens: int = MAX_NEW_TOKENS) -> Any:
    raw = generate_text(messages, max_new_tokens=max_new_tokens)
    try:
        return extract_json(raw)
    except Exception:
        repair_messages = [
            {"role": "system", "content": "Return valid JSON only. Do not add explanations, markdown, or code fences."},
            {
                "role": "user",
                "content": (
                    "The following model output is intended to be JSON but is invalid. "
                    "Fix only the JSON syntax and preserve the data as much as possible.\n\n"
                    f"{raw}"
                ),
            },
        ]
        return extract_json(generate_text(repair_messages, max_new_tokens=max_new_tokens))


def embed_texts(texts: list[str]) -> list[list[float]]:
    clean_texts = [str(text or "") for text in texts]
    if not clean_texts:
        return []
    vectors: list[list[float]] = []
    client = _client()
    for start in range(0, len(clean_texts), EMBEDDING_BATCH_SIZE):
        batch = clean_texts[start : start + EMBEDDING_BATCH_SIZE]
        kwargs: dict[str, Any] = {"model": EMBEDDING_MODEL_ID, "input": batch}
        if EMBEDDING_DIMENSIONS_INT:
            kwargs["dimensions"] = EMBEDDING_DIMENSIONS_INT
        last_exc: Exception | None = None
        for attempt in range(1, OPENAI_EMBED_MAX_RETRIES + 1):
            try:
                response = client.embeddings.create(**kwargs)
                break
            except Exception as exc:
                last_exc = exc
                if attempt >= OPENAI_EMBED_MAX_RETRIES:
                    raise RuntimeError(
                        f"OpenAI embedding request failed after {OPENAI_EMBED_MAX_RETRIES} attempts: {exc}"
                    ) from exc
                time.sleep(min(OPENAI_RETRY_BASE_SEC * (2 ** (attempt - 1)), 30.0))
        else:
            raise RuntimeError(f"OpenAI embedding request failed: {last_exc}")
        vectors.extend(item.embedding for item in response.data)
    return vectors


def _object_to_dict(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if hasattr(value, "model_dump"):
        return value.model_dump()
    if hasattr(value, "dict"):
        return value.dict()
    if hasattr(value, "__dict__"):
        return dict(value.__dict__)
    return {}


def _format_timestamp(seconds: float) -> str:
    seconds = max(0, int(float(seconds or 0)))
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def transcribe_audio_file(
    audio_path: str,
    *,
    language: str | None = "ko",
    batch_size: int = 16,
    align: bool = True,
    diarize: bool = True,
    hf_token: str | None = None,
) -> dict[str, Any]:
    _ = (batch_size, align, hf_token)
    load_stt_model()
    with open(audio_path, "rb") as audio_file:
        kwargs: dict[str, Any] = {"file": audio_file, "model": STT_MODEL_ID}
        if language:
            kwargs["language"] = language
        if diarize and "diarize" in STT_MODEL_ID:
            kwargs["response_format"] = "diarized_json"
            kwargs["chunking_strategy"] = "auto"
        else:
            kwargs["response_format"] = "verbose_json"
        transcript = _client().audio.transcriptions.create(**kwargs)

    payload = _object_to_dict(transcript)
    raw_segments = payload.get("segments") or []
    normalized_segments: list[dict[str, Any]] = []
    for index, item in enumerate(raw_segments):
        seg = _object_to_dict(item)
        text = str(seg.get("text") or seg.get("transcript") or "").strip()
        if not text:
            continue
        speaker = (
            seg.get("speaker")
            or seg.get("speaker_id")
            or seg.get("speaker_label")
            or seg.get("label")
            or f"SPEAKER_{index:02d}"
        )
        start = float(seg.get("start") or seg.get("start_time") or 0)
        normalized_segments.append(
            {
                **seg,
                "speaker": str(speaker),
                "start": start,
                "time": f"[{_format_timestamp(start)}]",
                "text": text,
                "content": text,
            }
        )
    text = "\n".join(
        f"{segment['speaker']}: {segment['text']}"
        for segment in normalized_segments
    ).strip()
    if not text:
        text = str(payload.get("text") or "").strip()
    return {
        "text": text,
        "language": payload.get("language") or language,
        "segments": normalized_segments,
        "utterances": normalized_segments,
        "device": "api",
        "compute_type": "api",
    }


def _decode_text(content: bytes) -> tuple[str, str]:
    last_error: Exception | None = None
    for encoding in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return content.decode(encoding), encoding
        except UnicodeDecodeError as exc:
            last_error = exc
    raise ValueError(f"Unable to decode text file: {last_error}")


def _extract_docx_text(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    lines = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines).strip()


def _extract_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[page {index}]\n{text}")
    return "\n\n".join(pages).strip()


def _content_type(filename: str, explicit: str = "") -> str:
    if explicit:
        return explicit
    guessed, _ = mimetypes.guess_type(filename)
    return guessed or "application/octet-stream"


def _image_data_url(image_bytes: bytes, filename: str, content_type: str = "") -> str:
    mime = _content_type(filename, content_type)
    if not mime.startswith("image/"):
        mime = "image/png"
    encoded = base64.b64encode(image_bytes).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def _ocr_image_bytes(image_bytes: bytes, filename: str, content_type: str = "") -> str:
    response = _client().chat.completions.create(
        model=OCR_MODEL_ID,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "이미지 안의 모든 읽을 수 있는 텍스트를 한국어 기준으로 정확히 추출하세요. "
                            "표는 행 단위로 보존하고, 설명 없이 추출 텍스트만 반환하세요."
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": _image_data_url(image_bytes, filename, content_type)}},
                ],
            }
        ],
        temperature=0,
        max_tokens=MAX_NEW_TOKENS,
    )
    return str(response.choices[0].message.content or "").strip()


def _render_pdf_pages(content: bytes) -> list[tuple[str, bytes]]:
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    matrix = fitz.Matrix(OCR_RENDER_SCALE, OCR_RENDER_SCALE)
    pages: list[tuple[str, bytes]] = []
    for page_index in range(min(OCR_MAX_PDF_PAGES, len(doc))):
        pixmap = doc.load_page(page_index).get_pixmap(matrix=matrix, alpha=False)
        pages.append((f"page-{page_index + 1}.png", pixmap.tobytes("png")))
    return pages


def ocr_file_bytes(file_bytes: bytes, filename: str, content_type: str = "") -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    mime = _content_type(filename, content_type)
    if ext == ".txt" or mime.startswith("text/"):
        return _decode_text(file_bytes)[0].strip()
    if ext == ".docx":
        return _extract_docx_text(file_bytes)
    if ext == ".pdf":
        direct_text = ""
        if OCR_PDF_TEXT_FIRST:
            try:
                direct_text = _extract_pdf_text(file_bytes)
            except Exception:
                direct_text = ""
        if len(direct_text.strip()) >= OCR_MIN_DIRECT_TEXT_CHARS:
            return direct_text.strip()
        page_texts = []
        for page_name, image_bytes in _render_pdf_pages(file_bytes):
            text = _ocr_image_bytes(image_bytes, page_name, "image/png")
            if text:
                page_texts.append(f"[{page_name}]\n{text}")
        return "\n\n".join(page_texts).strip()
    if mime.startswith("image/") or ext in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}:
        return _ocr_image_bytes(file_bytes, filename, content_type)
    raise ValueError(f"Unsupported OCR file type: {filename}")
